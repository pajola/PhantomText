from ..attack_base_injection import AttackBase
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader, PdfWriter
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
from reportlab.lib.pagesizes import A4
import random
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class OutOfBoundInjection(AttackBase):
    def __init__(self, modality="default", file_format="pdf"):
        if file_format == "pdf" or file_format == "html" or file_format == "docx":
            self.file_format = file_format

            if modality == "default" or modality == 'specific-position':
                self.modality = modality
            else:
                raise ValueError(f"Modality {modality} is not valid. Valid modality is 'default' (random out of bound position) or 'specific-position' (injection at a specific position)")
        else:
            raise ValueError(f"File format {file_format} is not valid. Valid file format is 'pdf', 'docx' or 'html'")
       

    def apply(self, input_document, injection, font_size=12, x_coord=100, y_coord=730, image_file=None, output_path=None):
        if not input_document:
            raise ValueError("Input document is required")
        
        # Check if injection exists
        if not injection:
            raise ValueError("Injection is required")
        
        if self.modality == 'specific-position' and (not x_coord or not y_coord):
            raise ValueError("x and y coordinates must be provided for specific position modality")

        # Check if output path is provided
        if not output_path:
            output_path = f"{input_document.split('.')[0]}_injected.{self.file_format}"
        
        if self.file_format == "pdf":

            buffer = BytesIO()
            c = canvas.Canvas(buffer)    

            if self.modality == 'specific-position':
                x, y = x_coord, y_coord
            else:
                x,y=self.get_random_out_of_bounds_position_pdf(c.stringWidth(injection,fontSize=font_size), font_size)

            c.drawString(x, y, injection)  # Adjust coordinates as needed

            c.save()
            buffer.seek(0)

            # Read the existing PDF
            existing_pdf = PdfReader(input_document)
            new_pdf = PdfReader(buffer)

            # Create a PdfWriter object to combine PDFs
            pdf_writer = PdfWriter()

            # Merge the new content with the existing PDF
            if len(existing_pdf.pages) > 0:
                existing_page = existing_pdf.pages[0]  # Assuming there's only one page
                new_page = new_pdf.pages[0]  # We only have one page from ReportLab
                
                existing_page.merge_page(new_page)  # Merge new content onto existing page
                
                pdf_writer.add_page(existing_page)

            # Write out the combined PDF to a file
            with open(output_path, 'wb') as f:
                pdf_writer.write(f)        

        elif self.file_format == "docx":

            doc = Document(input_document)

            section = doc.sections[0]
            if self.modality == 'specific-position':
                x, y = x_coord, y_coord
            else:
                x= 21
                y=0

            p = doc.add_paragraph()
            run = p.add_run(injection)
            p_format = p.paragraph_format
            p_format.left_indent = Inches(x)
            p_format.space_before = Inches(y)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Save the modified DOCX to the output path
            doc.save(output_path)

        elif self.file_format == "html":
            soup = BeautifulSoup(open(input_document), 'html.parser')
            # Create a new tag for the injection
            new_tag = soup.new_tag('div')
            
            # Set the text content)
            new_tag.string = injection

            if self.modality == 'specific-position':
                x= f"{x_coord}px"
                y= f"{y_coord}px"
            else:
                position = self.get_random_out_of_bounds_position_html()
                x= position['left']

            new_tag['style'] = f"position: absolute; left: {x}; top: {-2500};"
            soup.body.append(new_tag)
            
            with open(output_path, 'w') as f:
                f.write(str(soup))
            
    def get_random_out_of_bounds_position_html(self):
        """
        Generate random out-of-bounds positions based on predefined zones.
        Returns a dictionary containing the 'top' and 'left' CSS values.
        """
        # Define the zones for out-of-bounds positions
        zones = {
            "left": {"top": f"{random.uniform(0, 1000)}px", "left": f"{random.uniform(-2500, -2000)}px"},
            "top_left": {"top": f"{random.uniform(-2500, -2000)}px", "left": f"{random.uniform(-2500, -2000)}px"},
            "bottom_right": {"top": f"{random.uniform(2000, 2500)}px", "left": f"{random.uniform(-2000, -2500)}px"},
        }
        
        # Randomly choose a zone
        random_zone = random.choice(list(zones.keys()))
        return zones[random_zone]
    
    def get_random_out_of_bounds_position_pdf(self, text_width, font_size):
      
        # Default A4 size used by Canvas
        """
        Returns a random out-of-bounds position for a PDF injection.
        Coordinates are relative to the bottom-left corner of the page.
        The position is chosen randomly from one of 8 zones:
        top, bottom, left, right, top_left, top_right, bottom_left, bottom_right.
        The zones are defined as follows:
        - Edges (top, bottom, left, right): random position along the edge
        - Corners (top_left, top_right, bottom_left, bottom_right): random position
            within a square of size 300x300 centered at the corner
        The BUFFER constant is used to ensure that the injected text is not
        partially visible along the edge of the page.
        """
        PAGE_WIDTH, PAGE_HEIGHT = A4  # 595 x 842 points
        BUFFER = 30  # Extra padding to avoid partial visibility 
        
        # Choose a random zone (e.g., 'left', 'top_right', 'bottom_center', etc.)

        # zone = random.choice(zones)

        
        # Define all 8 zones with randomized coordinates
        zones = {
            # Edges
            "top": (
                random.uniform(0, PAGE_WIDTH), 
                random.uniform(PAGE_HEIGHT + BUFFER, PAGE_HEIGHT + 300)
            ),
            "bottom": (
                random.uniform(0, PAGE_WIDTH), 
                random.uniform(-300, -font_size-BUFFER)
            ),
            "left": (
                random.uniform(-300, -text_width - BUFFER), 
                random.uniform(0, PAGE_HEIGHT)
            ),
            "right": (
                random.uniform(PAGE_WIDTH + BUFFER, PAGE_WIDTH + 300), 
                random.uniform(0, PAGE_HEIGHT)
            ),
            # Corners
            "top_left": (
                random.uniform(-300, -text_width - BUFFER), 
                random.uniform(PAGE_HEIGHT + BUFFER, PAGE_HEIGHT + 300)
            ),
            "top_right": (
                random.uniform(PAGE_WIDTH + BUFFER, PAGE_WIDTH + 300), 
                random.uniform(PAGE_HEIGHT + BUFFER, PAGE_HEIGHT + 300)
            ),
            "bottom_left": (
                random.uniform(-300, -text_width - BUFFER), 
                random.uniform(-300, -font_size-BUFFER)
            ),
            "bottom_right": (
                random.uniform(PAGE_WIDTH + BUFFER, PAGE_WIDTH + 300), 
                random.uniform(-300,- font_size-BUFFER)
            ),
        }
        
        return zones[random.choice(list(zones.keys()))]

    def check(self, input_document):
        return True